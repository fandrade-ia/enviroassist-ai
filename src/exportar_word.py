"""EnviroAssist AI — Exportación Word"""
from pathlib import Path
from datetime import datetime
from io import BytesIO
import re

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

COLOR_VERDE = RGBColor(0x2E, 0x7D, 0x32)
COLOR_GRIS  = RGBColor(0x42, 0x42, 0x42)
COLOR_O     = RGBColor(0xEF, 0x6C, 0x00)

def _cfg(doc):
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

def _portada(doc, proyecto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("🌿 EnviroAssist AI")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = COLOR_VERDE
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(proyecto.get("nombre","Proyecto"))
    r2.bold = True; r2.font.size = Pt(14)
    doc.add_paragraph()
    t = doc.add_table(rows=4, cols=2); t.style = "Table Grid"
    for i,(k,v) in enumerate([
        ("Ubicación", proyecto.get("ubicacion","")),
        ("CCAA", proyecto.get("ccaa","")),
        ("Tipo", proyecto.get("tipo","")),
        ("Fecha", datetime.now().strftime("%d/%m/%Y")),
    ]):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v
        if t.rows[i].cells[0].paragraphs[0].runs:
            t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("⚠️  Borrador generado con IA. Revisar antes de presentación oficial.")
    r3.italic = True; r3.font.size = Pt(9); r3.font.color.rgb = COLOR_O
    doc.add_page_break()

def _texto(p, txt):
    txt = re.sub(r"[⚠️🔴🟠🟡🟢✅❌🌿🦅📂📋🚨]", "", txt)
    partes = re.split(r"(\*\*.*?\*\*|\*.*?\*)", txt)
    for parte in partes:
        if parte.startswith("**") and parte.endswith("**"):
            p.add_run(parte[2:-2]).bold = True
        elif parte.startswith("*") and parte.endswith("*"):
            p.add_run(parte[1:-1]).italic = True
        else:
            p.add_run(parte)

def _md_a_word(doc, texto):
    for linea in texto.split("\n"):
        if linea.startswith("## "):
            h = doc.add_heading(linea[3:].strip(), level=2)
            if h.runs: h.runs[0].font.color.rgb = COLOR_VERDE
        elif linea.startswith("# "):
            h = doc.add_heading(linea[2:].strip(), level=1)
            if h.runs: h.runs[0].font.color.rgb = COLOR_VERDE
        elif linea.startswith("### "):
            doc.add_heading(linea[4:].strip(), level=3)
        elif linea.startswith("- ") or linea.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _texto(p, linea[2:])
        elif linea.strip() in ("---","___"):
            doc.add_paragraph()
        elif "|" in linea and linea.strip().startswith("|"):
            pass  # tablas markdown se omiten (se muestran como texto)
        elif linea.strip() == "":
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            _texto(p, linea)

def markdown_a_docx_bytes(texto: str, proyecto: dict) -> bytes:
    doc = Document()
    _cfg(doc)
    _portada(doc, proyecto)
    _md_a_word(doc, texto)
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(f"EnviroAssist AI · Borrador · {datetime.now().strftime('%d/%m/%Y')}")
    r.font.size = Pt(8); r.italic = True
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
