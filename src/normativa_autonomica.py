"""
EnviroAssist AI — Normativa Autonómica de Especies Amenazadas
- LESRPE/CEEA: cargado desde data/normativa/lesrpe_2025.json
- Todas las CCAA (incluida Madrid): se procesan igual, a partir del catálogo
  que sube el usuario en Excel (.xlsx/.xls), CSV, PDF o Word (.docx)
"""

import json, io, re
import pandas as pd
from pathlib import Path
from openai import OpenAI
from dotenv import load_dotenv
load_dotenv()

SRC_DIR     = Path(__file__).parent
DATA_DIR    = SRC_DIR.parent / "data"
LESRPE_JSON = DATA_DIR / "normativa" / "lesrpe_2025.json"

LLM_MODEL     = "microsoft/phi-4"
LM_STUDIO_URL = "http://localhost:1234/v1"
client = OpenAI(base_url=LM_STUDIO_URL, api_key="lm-studio")

CCAA_LISTA = [
    "Andalucía","Aragón","Asturias","Baleares","Canarias",
    "Cantabria","Castilla-La Mancha","Castilla y León","Cataluña",
    "Extremadura","Galicia","La Rioja","Madrid","Murcia",
    "Navarra","País Vasco","Valencia"
]

CATEGORIAS_DEFAULT = {
    "EP":  {"nombre":"En peligro de extinción",             "emoji":"🟠","nivel":1},
    "E":   {"nombre":"En peligro de extinción",             "emoji":"🟠","nivel":1},
    "SAH": {"nombre":"Sensible a la alteración del hábitat","emoji":"🟢","nivel":2},
    "V":   {"nombre":"Vulnerable",                          "emoji":"🟡","nivel":3},
    "IE":  {"nombre":"De interés especial",                 "emoji":"🟢","nivel":4},
}

# Los códigos internos (EP/SAH/V/IE/E) se siguen usando para ordenar por
# nivel de amenaza, pero de cara al color se agrupan en solo 3 categorías:
# EP = naranja, V = amarillo, cualquier otra cosa (SAH, IE, LESRPE, "listado"...)
# = "Otro" en verde. Así el color no depende de cuántas categorías use cada CCAA.
COLOR_INFO = {
    "EP":   {"emoji":"🟠", "fill":"FFEDD5", "font":"9A3412", "label":"En peligro de extinción"},
    "V":    {"emoji":"🟡", "fill":"FEF9C3", "font":"854D0E", "label":"Vulnerable"},
    "OTRO": {"emoji":"🟢", "fill":"DCFCE7", "font":"166534", "label":"Otro"},
}


def _color_cod(cat_cod: str) -> str:
    """Agrupa cualquier código de categoría autonómica en los 3 colores del
    semáforo: EP y E -> naranja, V -> amarillo, todo lo demás -> verde ('Otro')."""
    if cat_cod in ("EP", "E"):
        return "EP"
    if cat_cod == "V":
        return "V"
    return "OTRO"


def _limpiar_prefijo_categoria(texto: str) -> str:
    """Quita prefijos de listado tipo 'A) ', 'b)', '1) ', 'a.' que traen
    algunos catálogos autonómicos porque copian el literal del articulado
    (p.ej. Decreto 18/1992 de Madrid: 'A) En peligro de extinción',
    'B) Sensibles a la alteración de su hábitat'...). Se aplica a
    cualquier CCAA, no solo a Madrid, ya que el prefijo es un artefacto
    del texto legal de origen, no información de la categoría en sí."""
    if not texto:
        return texto
    return re.sub(r"^\s*[A-Za-z0-9]+[\)\.]\s*", "", texto).strip()



def extraer_texto_pdf(archivo_bytes: bytes) -> str:
    """Extrae el texto de un PDF respetando el orden VISUAL de columnas.

    Los catálogos autonómicos suelen maquetarse a 1, 2 o más columnas (y a
    veces con varias tablas independientes una junto a otra en la misma
    página, como el Decreto 18/1992 de Madrid). Si se usa la extracción de
    texto por defecto (p.ej. pypdf), el orden de las palabras en el stream
    interno del PDF no siempre coincide con el orden visual, y eso hace que
    cabeceras de categoría y especies de columnas/tablas distintas se
    mezclen — la causa real de que species terminen con la categoría de
    otra tabla.

    Esta función, en cambio: (1) agrupa las palabras de cada página por
    posición X, detectando huecos horizontales amplios como límites entre
    columnas — sin asumir un número fijo de columnas, así funciona igual de
    bien en documentos a 1, 2, 3 o 4 columnas; (2) procesa cada columna de
    arriba a abajo, de izquierda a derecha; (3) dentro de cada columna,
    agrupa palabras por su coordenada Y para reconstruir líneas — lo que de
    paso junta en una misma línea pares "nombre común | nombre científico"
    de tablas a dos columnas, facilitando emparejarlos después."""
    import pdfplumber
    texto_paginas = []
    with pdfplumber.open(io.BytesIO(archivo_bytes)) as pdf:
        for page in pdf.pages:
            try:
                words = page.extract_words(use_text_flow=False, keep_blank_chars=False)
            except Exception:
                words = []
            if not words:
                continue
            xs = sorted(set(round(w["x0"], 1) for w in words))
            cortes = [(xs[i] + xs[i + 1]) / 2 for i in range(len(xs) - 1) if xs[i + 1] - xs[i] > 18]
            bordes = [0.0] + cortes + [page.width + 1]
            lineas_pagina = []
            for i in range(len(bordes) - 1):
                banda = [w for w in words if bordes[i] <= w["x0"] < bordes[i + 1]]
                if not banda:
                    continue
                filas = {}
                for w in banda:
                    filas.setdefault(round(w["top"]), []).append(w)
                for top in sorted(filas):
                    palabras_fila = sorted(filas[top], key=lambda w: w["x0"])
                    lineas_pagina.append(" ".join(w["text"] for w in palabras_fila))
            texto_paginas.append("\n".join(lineas_pagina))
    return "\n".join(texto_paginas)


def cargar_lesrpe() -> dict:
    if LESRPE_JSON.exists():
        with open(LESRPE_JSON, encoding="utf-8") as f:
            return json.load(f)
    return {}


def get_normativa_inicial() -> dict:
    """No hay ninguna CCAA precargada: todas, incluida Madrid, se cargan de la
    misma forma — el usuario sube su catálogo (Excel, CSV, PDF o Word) desde
    la pestaña 'Normativa autonómica' y se procesa igual para todas."""
    return {}


_COL_ESPECIE_KEYS = ["cientifico", "taxon", "especie", "nombre", "species"]


def _promover_cabecera(df: pd.DataFrame) -> pd.DataFrame:
    """Algunos catálogos autonómicos (p.ej. Andalucía) tienen una fila en
    blanco o un título con celdas combinadas antes de la cabecera real de la
    tabla. pandas coge por defecto la primera fila como cabecera, lo que deja
    columnas 'Unnamed: 0', 'Unnamed: 1'... y _parsear_hoja no encuentra nada.
    Esta función busca, entre las primeras filas, la que realmente contiene
    los nombres de columna (p.ej. 'NOMBRE CIENTIFICO') y la promueve como
    cabecera, descartando lo que hubiera por encima."""
    max_filas = min(10, len(df))
    for i in range(max_filas):
        fila = df.iloc[i].astype(str).str.lower()
        if any(fila.str.contains(k, na=False).any() for k in _COL_ESPECIE_KEYS):
            nuevo = df.iloc[i + 1:].copy()
            nuevo.columns = [str(c).strip() for c in df.iloc[i]]
            return nuevo.reset_index(drop=True)
    return df


def _parsear_hoja(df: pd.DataFrame) -> dict:
    df = _promover_cabecera(df)
    df = df.copy()
    df.columns = [str(c).strip().lower() for c in df.columns]
    col_esp = next((c for c in df.columns if any(k in c for k in _COL_ESPECIE_KEYS)), None)
    if col_esp is None:
        return {}
    col_cat = next((c for c in df.columns if any(k in c for k in
        ["categoria", "estatus", "estado", "cat", "proteccion", "amenaza", "status"])), None)
    especies = {}
    for _, row in df.iterrows():
        nombre_sp = str(row[col_esp]).strip() if pd.notna(row[col_esp]) else ""
        if not nombre_sp or nombre_sp.lower() in ["nan", "none", ""]:
            continue
        cat = str(row[col_cat]).strip() if col_cat and pd.notna(row[col_cat]) else "IE"
        cat_norm = _normalizar_categoria(cat)
        if cat_norm is None:
            continue  # p.ej. "EX" (extinta): no se incluye como amenaza activa
        grupo = next((str(row[c]).strip() for c in df.columns
                      if any(k in c for k in ["grupo", "clase", "orden"]) and pd.notna(row[c])), "No especificado")
        nombre_comun = next((str(row[c]).strip() for c in df.columns
                             if any(k in c for k in ["comun", "vulgar", "common"]) and pd.notna(row[c])), "")
        partes = nombre_sp.split()
        if len(partes) >= 2:
            clave = partes[0] + " " + partes[1]
            especies[clave] = {"cat": cat_norm, "cat_original": cat, "grupo": grupo, "nombre_comun": nombre_comun}
    return especies


def procesar_excel_normativa(archivo_bytes: bytes, ccaa: str, nombre: str) -> dict:
    """Soporta libros Excel de cualquier tamaño y con varias hojas (algunas
    CCAA separan fauna y flora en pestañas distintas del mismo archivo): se
    procesan todas y se fusionan las especies encontradas.

    Soporta tanto .xlsx (OOXML, engine openpyxl) como el formato binario
    antiguo .xls de Excel 97-2003 (engine xlrd), que openpyxl NO puede leer.
    Si la extensión no es concluyente, se prueban ambos engines.

    Se lee cada hoja SIN asumir que la fila 1 es la cabecera (header=None):
    _parsear_hoja/_promover_cabecera detectan la fila de cabecera real,
    porque varios catálogos autonómicos llevan una fila de título o celdas
    combinadas por encima de la tabla."""
    especies = {}
    ext = Path(nombre).suffix.lower()
    engines = ["xlrd", "openpyxl"] if ext == ".xls" else ["openpyxl", "xlrd"]

    hojas = None
    for engine in engines:
        try:
            hojas = pd.read_excel(io.BytesIO(archivo_bytes), engine=engine,
                                  sheet_name=None, header=None)
            break
        except Exception:
            continue

    if hojas is not None:
        for df in hojas.values():
            especies.update(_parsear_hoja(df))
    else:
        df = None
        for enc in ["utf-8-sig", "latin-1", "cp1252"]:
            for sep in [";", ",", "\t"]:
                try:
                    df = pd.read_csv(io.BytesIO(archivo_bytes), encoding=enc,
                                     sep=sep, on_bad_lines="skip")
                    if len(df.columns) > 1:
                        break
                except Exception:
                    continue
            if df is not None:
                break
        if df is not None:
            especies = _parsear_hoja(df)

    if not especies:
        return {}
    return {"ccaa": ccaa, "decreto": f"Cargado desde: {nombre}", "categorias": CATEGORIAS_DEFAULT, "especies": especies}


def procesar_docx_normativa(archivo_bytes: bytes, ccaa: str, nombre: str) -> dict:
    """Procesa un catálogo autonómico en Word (.docx). Funciona igual para
    cualquier CCAA (incluida Madrid, que ya no lleva datos hardcodeados):

    1) Si el documento tiene tablas (caso típico de un anexo con columnas
       Grupo / Nombre científico / Categoría...), cada tabla se procesa
       exactamente igual que una hoja Excel: se reutiliza _parsear_hoja, que
       detecta la fila de cabecera real y localiza las columnas de especie
       y categoría por palabras clave.
    2) Si no hay tablas o no se extrae nada de ellas (catálogo redactado
       como texto corrido, con cabeceras de categoría seguidas de nombres
       científicos), se cae al mismo motor de patrones de texto que se usa
       para PDF, sobre el texto de los párrafos del documento.

    En ambos casos se conserva el texto literal de la categoría tal como
    aparece en el documento (columna de la tabla, o cabecera de sección)."""
    from docx import Document
    doc = Document(io.BytesIO(archivo_bytes))

    especies = {}
    for tabla in doc.tables:
        filas = [[c.text.strip() for c in fila.cells] for fila in tabla.rows]
        if not filas:
            continue
        df = pd.DataFrame(filas)
        especies.update(_parsear_hoja(df))

    if especies:
        texto_parrafos = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        decreto = _extraer_decreto_regex(texto_parrafos)
        return {"ccaa": ccaa, "decreto": decreto or f"Cargado desde: {nombre}",
                "categorias": CATEGORIAS_DEFAULT, "especies": especies}

    # Sin tablas útiles: tratar como texto corrido, igual que un PDF
    texto_parrafos = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return procesar_pdf_normativa(texto_parrafos, ccaa)


def procesar_pdf_normativa(texto_pdf: str, ccaa: str) -> dict:
    """Extrae especies y decreto de un PDF de normativa autonómica, sea cual
    sea su tamaño (desde un decreto de 6 páginas hasta una norma consolidada
    de 80+ páginas).

    MÉTODO PRINCIPAL — extracción por patrones de texto (determinista):
    los catálogos autonómicos de especies amenazadas casi siempre siguen la
    misma estructura: cabeceras de categoría ("EN PELIGRO DE EXTINCIÓN",
    "VULNERABLES"...), cabeceras de grupo taxonómico (AVES, Fam. XXXAE...) y,
    debajo, nombres científicos en líneas propias. Al ser regex sobre texto
    ya extraído, procesa el documento completo (sin recortes) en milisegundos
    y no depende de que un modelo local genere JSON válido — que es
    precisamente lo que fallaba con documentos largos: los anexos con
    cientos de especies superaban el límite de tokens de salida de phi-4 y
    el JSON quedaba cortado a medias.

    MÉTODO DE RESERVA — solo si el patrón anterior no encuentra ninguna
    especie (documento con formato atípico): se trocea el texto completo en
    fragmentos y se procesa cada uno con phi-4, sin límite de fragmentos, así
    que también escala con el tamaño del documento.
    """
    decreto = _extraer_decreto_regex(texto_pdf)
    especies = _extraer_especies_regex(texto_pdf)
    metodo = "patrones de texto"

    if not especies:
        especies, decreto_llm = _extraer_especies_llm(texto_pdf, ccaa)
        decreto = decreto or decreto_llm
        metodo = "phi-4 (reserva)"

    return {
        "ccaa": ccaa,
        "decreto": decreto or "No identificado",
        "categorias": CATEGORIAS_DEFAULT,
        "especies": especies,
        "metodo_extraccion": metodo,
    }


_CAT_HEADERS = [
    (re.compile(r'peligro\s+de\s+extinci[oó]n', re.I), "EP"),
    (re.compile(r'sensible', re.I), "SAH"),
    (re.compile(r'vulnerabl', re.I), "V"),
    (re.compile(r'inter[eé]s\s+especial', re.I), "IE"),
    (re.compile(r'extint', re.I), None),  # extintas: no se añaden como amenazadas activas
]

_GRUPO_HEADERS = [
    (re.compile(r'^\**\s*aves\s*\**$', re.I), "Aves"),
    (re.compile(r'^\**\s*mam[ií]feros\s*\**$', re.I), "Mamíferos"),
    (re.compile(r'^\**\s*reptiles\s*\**$', re.I), "Reptiles"),
    (re.compile(r'^\**\s*anfibios\s*\**$', re.I), "Anfibios"),
    (re.compile(r'^\**\s*peces\s*\**$', re.I), "Peces"),
    (re.compile(r'^\**\s*insectos\s*\**$', re.I), "Insectos"),
    (re.compile(r'^\**\s*crust[aá]ceos\s*\**$', re.I), "Crustáceos"),
    (re.compile(r'^\**\s*gaster[oó]podos\s*\**$', re.I), "Gasterópodos"),
    (re.compile(r'^\**\s*bivalvos\s*\**$', re.I), "Bivalvos"),
    (re.compile(r'^\**\s*moluscos\s*\**$', re.I), "Moluscos"),
    (re.compile(r'^\**\s*invertebrados\s*\**$', re.I), "Invertebrados"),
    (re.compile(r'^\**\s*ar[aá]cnidos\s*\**$', re.I), "Arácnidos"),
    (re.compile(r'^\**\s*hongos\s*\**$', re.I), "Hongos"),
    (re.compile(r'especies\s+de\s+flora', re.I), "Flora"),
    (re.compile(r'especies\s+de\s+fauna', re.I), "Fauna"),
]

# Género especie [subsp./var. epíteto], sin anclar al inicio de línea: en
# tablas a dos columnas (nombre común | nombre científico) unidas en una
# misma línea por extraer_texto_pdf, el nombre científico no siempre empieza
# en la posición 0.
_GENUS_EPITHET_RE = re.compile(
    r'([A-ZÀ-Ý][a-zà-ÿ]{2,}(?:\s*\([A-ZÀ-Ý][a-zà-ÿ]+\))?\s+[a-zà-ÿ]{3,}'
    r'(?:\s+(?:subsp\.|var\.|f\.)\s*[a-zà-ÿ\-]+)?)'
)

# Lo que viene DESPUÉS de un candidato a nombre científico, para confirmar
# que de verdad lo es y no es, p.ej., un nombre común en español que por
# casualidad tiene la forma "Palabra Capitalizada + palabra en minúscula"
# (como "Cigüeña negra"): en el 99% de los catálogos oficiales, el nombre
# científico va seguido inmediatamente de una cita de autor — un apellido
# capitalizado, opcionalmente con más autores unidos por "&"/"et"/"ex", una
# coma y un año, o el nombre de la familia/autor entre paréntesis. Esta
# señal es válida tanto para el formato de Madrid ("Linnaeus") como el de
# Andalucía ("Cobo, 1962", "R. Galán, Baral & A. Ortega").
_AUTOR_RE = re.compile(
    r'^\s*(?:\([^()]{2,90}\)\.?'
    r'|[A-ZÀ-Ý][\wÀ-ÿ.\'’\-]{1,30}(?:[\s,]+(?:&|et|ex)\s*[A-ZÀ-Ý]?[\wÀ-ÿ.\'’\-]{0,30})*\s*,?\s*\d{0,4}\)?)'
)

# Palabras que, aunque encajen con el patrón "Capitalizada + minúscula", casi
# nunca son un género o epíteto real: preposiciones, artículos y vocabulario
# habitual del articulado legal español. No pretende ser exhaustiva —
# combinada con la exigencia de cita de autor de arriba, basta para eliminar
# la inmensa mayoría de falsos positivos del texto corrido del decreto.
_PALABRAS_VACIAS = {
    "de", "la", "el", "las", "los", "en", "por", "para", "con", "del", "al",
    "su", "sus", "que", "un", "una", "uno", "no", "si", "es", "son", "como",
    "este", "esta", "estos", "estas", "dicho", "dicha", "dichos", "dichas",
    "ver", "vease", "véase", "pagina", "página", "boletin", "boletín",
    "dado", "dentro", "desde", "sobre", "entre", "cada", "todo", "toda",
    "todos", "todas", "otro", "otra", "sin", "mas", "más", "regional",
    "nacional", "catalogo", "catálogo", "especial", "general", "especies",
    "articulo", "artículo", "disposición", "disposicion", "capítulo",
    "capitulo", "sección", "seccion", "anexo", "decreto", "ley",
}


def _candidato_valido(nombre: str) -> bool:
    partes = nombre.replace("(", " ").replace(")", " ").split()
    if len(partes) < 2:
        return False
    genero, epiteto = partes[0].lower(), partes[1].lower()
    return genero not in _PALABRAS_VACIAS and epiteto not in _PALABRAS_VACIAS


def _candidatos_especie_en_linea(linea: str) -> list[tuple[str, str, str]]:
    """Devuelve una lista de (nombre_cientifico, nombre_comun_previo,
    cita_autor) para cada nombre científico verificado que aparezca en la
    línea. Puede haber más de uno si la línea combina nombre común +
    nombre científico (tablas de dos columnas fusionadas en una línea) o
    varias especies seguidas."""
    candidatos = [m for m in _GENUS_EPITHET_RE.finditer(linea) if _candidato_valido(m.group(1))]

    # Si dos candidatos son adyacentes (sin texto real entre ellos), el
    # primero es casi siempre el nombre común de la especie del segundo
    # candidato (p.ej. "Sisón" + "Tetrax tetrax" -> "Sisón Tetrax tetrax"
    # se detecta primero como candidato "Sisón Tetrax", que hay que
    # descartar en favor de "Tetrax tetrax").
    filtrados = []
    for i, m in enumerate(candidatos):
        if i + 1 < len(candidatos):
            hueco = linea[m.end():candidatos[i + 1].start()]
            if hueco.strip(" ,") == "":
                continue
        filtrados.append(m)

    resultados = []
    for m in filtrados:
        resto = linea[m.end():]
        autor = _AUTOR_RE.match(resto)
        if not autor or len(autor.group().strip()) < 2:
            continue
        nombre = re.sub(r"\s+", " ", m.group(1)).strip()
        antes = linea[:m.start()].strip(" -•\t,;")
        # el "antes" solo es un nombre común plausible si es corto y no
        # contiene ya otro nombre científico (evita arrastrar texto de
        # líneas mal cortadas)
        nombre_comun = antes if antes and len(antes) <= 45 and not _GENUS_EPITHET_RE.search(antes) else ""
        resultados.append((nombre, nombre_comun, autor.group().strip()))
    return resultados


def _extraer_decreto_regex(texto: str) -> str:
    m = re.search(r'(Decret[oa]\s+(?:Foral\s+)?[\d/]+.{0,90})', texto[:6000], re.IGNORECASE)
    return m.group(1).strip() if m else ""


def _es_cabecera(linea: str) -> bool:
    """Detecta líneas de cabecera (categoría, grupo, familia) a descartar como especie."""
    l = linea.strip(" -•\t")
    if not l:
        return True
    if l.startswith("Fam.") or l.startswith("ANEXO") or l.startswith("CAPÍTULO"):
        return True
    letras = [c for c in l if c.isalpha()]
    # todo mayúsculas -> es una cabecera (categoría, sección...), sin límite
    # de palabras: algunas CCAA usan frases largas como cabecera de sección
    # ("ESPECIES QUE SE INCLUYEN EN LA CATEGORÍA 'EN PELIGRO DE EXTINCIÓN'").
    if letras and all(c.isupper() for c in letras) and len(l.split()) <= 20:
        return True
    return False


# Nota: la validación de "palabras que nunca son género/epíteto" vive ahora
# en _PALABRAS_VACIAS, usada por _candidato_valido().


def _extraer_especies_regex(texto: str) -> dict:
    """Recorre el texto completo línea a línea; escala linealmente con el
    tamaño del documento (sin límites de caracteres ni de páginas)."""
    cat_actual = None
    cat_original_actual = ""
    grupo_actual = "—"
    especies = {}
    for linea in texto.splitlines():
        stripped = linea.strip()
        if not stripped:
            continue

        limpio = stripped.strip(" -•\t")
        es_hdr = _es_cabecera(stripped)
        cabecera_detectada = False

        if es_hdr:
            for patron, code in _CAT_HEADERS:
                if patron.search(limpio):
                    cat_actual = code
                    cat_original_actual = limpio  # texto literal del decreto, p.ej. "EN PELIGRO DE EXTINCIÓN"
                    cabecera_detectada = True
                    break
        for patron, label in _GRUPO_HEADERS:
            if patron.search(limpio):
                grupo_actual = label
                cabecera_detectada = True
                break
        if cabecera_detectada or es_hdr:
            continue
        if cat_actual is None:
            continue  # aún no hemos entrado en ninguna sección de categoría

        for nombre, comun, _autor in _candidatos_especie_en_linea(stripped):
            if nombre not in especies:
                especies[nombre] = {"cat": cat_actual, "cat_original": cat_original_actual,
                                    "grupo": grupo_actual, "nombre_comun": comun}
    return especies


def _reparar_json_truncado(texto: str) -> str:
    """Si la respuesta del LLM se cortó a mitad de un valor (límite de
    tokens), intenta cerrar las estructuras abiertas para poder parsear lo
    que sí llegó completo, en vez de descartar todo el fragmento."""
    texto = texto.rstrip()
    # corta en la última coma de un par clave-valor completo si el final
    # está a medias (p.ej. una cadena sin cerrar)
    ultima_llave_cierre = texto.rfind("}")
    if ultima_llave_cierre != -1 and ultima_llave_cierre < len(texto) - 1:
        texto = texto[:ultima_llave_cierre + 1]
    abiertas_llave = texto.count("{") - texto.count("}")
    abiertas_corchete = texto.count("[") - texto.count("]")
    texto += "}" * max(abiertas_llave, 0)
    texto += "]" * max(abiertas_corchete, 0)
    return texto


def _extraer_especies_llm(texto_pdf: str, ccaa: str) -> tuple[dict, str]:
    """Reserva vía phi-4: trocea el texto completo (sin límite de
    fragmentos, escala con el tamaño del documento) y fusiona lo que cada
    fragmento devuelva, tolerando JSON truncado o inválido en fragmentos
    sueltos sin abortar el resto."""
    CHUNK_SIZE = 4000  # fragmentos más pequeños = menos riesgo de truncar la salida
    chunks = [texto_pdf[i:i + CHUNK_SIZE] for i in range(0, len(texto_pdf), CHUNK_SIZE)] or [""]
    decreto_final = ""
    todas_especies = {}
    for i, chunk in enumerate(chunks):
        if not chunk.strip():
            continue
        prompt = f"""Eres un experto en legislación ambiental española.
Este es el FRAGMENTO {i+1} de {len(chunks)} de un documento del catálogo de
especies amenazadas de {ccaa}. Extrae SOLO las especies de fauna o flora con
su categoría de amenaza que aparezcan explícitamente en este fragmento
(puede haber fragmentos sin ninguna especie; en ese caso "especies": {{}}).

TEXTO:
{chunk}

Responde ÚNICAMENTE con JSON válido (sin texto adicional, sin markdown):
{{
  "decreto": "nombre y fecha del decreto si aparece en este fragmento, si no, cadena vacía",
  "especies": {{
    "Nombre cientifico": {{"cat":"EP","grupo":"Aves","nombre_comun":"Nombre común"}}
  }}
}}
Usa: EP=En peligro de extinción, SAH=Sensible a la alteración del hábitat, V=Vulnerable, IE=Interés especial."""
        data = None
        try:
            resp = client.chat.completions.create(
                model=LLM_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1, max_tokens=2500,
            )
            texto = resp.choices[0].message.content.strip()
            match = re.search(r'\{.*\}', texto, re.DOTALL)
            if match:
                try:
                    data = json.loads(match.group())
                except json.JSONDecodeError:
                    try:
                        data = json.loads(_reparar_json_truncado(match.group()))
                    except Exception:
                        data = None
        except Exception as e:
            print(f"Error PDF (fragmento {i+1}/{len(chunks)}): {e}")
        if not data:
            continue
        if not decreto_final and data.get("decreto"):
            decreto_final = data["decreto"]
        for nombre_sp, info in (data.get("especies") or {}).items():
            if nombre_sp and nombre_sp not in todas_especies:
                todas_especies[nombre_sp] = info
    return todas_especies, decreto_final


def cruzar_con_normativa(especie: str, normativa: dict) -> dict | None:
    if not normativa or not normativa.get("especies"):
        return None
    partes = especie.strip().split()
    if len(partes) < 2:
        return None
    clave2 = partes[0] + " " + partes[1]
    especies = normativa["especies"]
    info = especies.get(clave2) or especies.get(especie.strip())
    if info is None:
        for k in especies:
            if k.lower().startswith(clave2.lower()):
                info = especies[k]; break
    if not info:
        return None
    cat_cod = info.get("cat","IE")
    cats = normativa.get("categorias", CATEGORIAS_DEFAULT)
    cat_info = cats.get(cat_cod, CATEGORIAS_DEFAULT.get(cat_cod, {}))
    color_cod = _color_cod(cat_cod)
    color_info = COLOR_INFO[color_cod]
    # Texto a mostrar: EXACTAMENTE como viene en la normativa de origen
    # (p.ej. "EN", "LISTADO", "VU" en Andalucía). Si no se conservó el
    # literal (p.ej. catálogo de Madrid precargado), se usa como respaldo
    # el nombre genérico de la categoría.
    cat_original = info.get("cat_original") or (
        cat_info.get("nombre", cat_cod) if isinstance(cat_info, dict) else str(cat_info)
    )
    cat_original = _limpiar_prefijo_categoria(cat_original)
    return {
        "cat_cod": cat_cod,
        "cat_nombre": cat_original,
        "color_cod": color_cod,
        "emoji": color_info["emoji"],
        "nivel": cat_info.get("nivel",5) if isinstance(cat_info,dict) else 5,
        "nombre_comun": info.get("nombre_comun",""),
        "decreto": normativa.get("decreto",""),
    }


def _normalizar_categoria(cat: str) -> str | None:
    c = cat.lower().strip()
    # quita anotaciones tipo "VU (Mediterráneo)" o "EN\nRíos Guadiana..."
    c = re.split(r'[\(\n]', c)[0].strip()
    if re.search(r'^ex$|\bextint', c):
        return None  # extinta: no se incluye como amenaza activa
    if re.search(r'\bpeligro\b|\bextinci[oó]n\b|\bextinction\b|^ep$|^pe$|^cr$|^en$', c):
        return "EP"
    if re.search(r'\bsensible\b|\balteraci[oó]n\b|\bh[aá]bitat\b|^sah$', c):
        return "SAH"
    if re.search(r'\bvulnerable\b|^vu$|^v$', c):
        return "V"
    return "IE"
